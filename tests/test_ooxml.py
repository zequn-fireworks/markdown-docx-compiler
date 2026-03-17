"""Tests for backend/docx/ooxml.py — low-level OOXML helpers."""

from __future__ import annotations

from docx import Document
from docx.shared import RGBColor

from markdown_docx_compiler.backend.docx.ooxml import (
    add_page_field,
    rgb,
    set_cell_shading,
    set_cell_vertical_alignment,
    set_cell_width,
    set_paragraph_bottom_border,
    set_paragraph_left_border,
    set_paragraph_shading,
    set_secondary_font,
    set_table_borders,
    set_table_cell_margins,
    set_table_grid,
    set_table_layout_fixed,
    set_table_width_dxa,
)


class TestRgb:
    def test_valid_hex(self) -> None:
        color = rgb("FF8000")
        assert color == RGBColor(0xFF, 0x80, 0x00)

    def test_hex_with_hash(self) -> None:
        color = rgb("#1F2937")
        assert color == RGBColor(0x1F, 0x29, 0x37)

    def test_black(self) -> None:
        assert rgb("000000") == RGBColor(0, 0, 0)

    def test_white(self) -> None:
        assert rgb("FFFFFF") == RGBColor(0xFF, 0xFF, 0xFF)

    def test_lowercase(self) -> None:
        assert rgb("ff0000") == RGBColor(0xFF, 0, 0)


def _make_doc_with_table(rows: int = 2, cols: int = 2) -> Document:
    doc = Document()
    doc.add_table(rows=rows, cols=cols)
    return doc


class TestTableHelpers:
    def test_set_table_layout_fixed(self) -> None:
        doc = _make_doc_with_table()
        table = doc.tables[0]
        set_table_layout_fixed(table)
        xml = table._tbl.xml
        assert 'w:type="fixed"' in xml

    def test_set_table_layout_fixed_idempotent(self) -> None:
        doc = _make_doc_with_table()
        table = doc.tables[0]
        set_table_layout_fixed(table)
        set_table_layout_fixed(table)
        assert table._tbl.xml.count('w:type="fixed"') == 1

    def test_set_table_width_dxa(self) -> None:
        doc = _make_doc_with_table()
        table = doc.tables[0]
        set_table_width_dxa(table, 9360)
        xml = table._tbl.xml
        assert 'w:w="9360"' in xml
        assert 'w:type="dxa"' in xml

    def test_set_table_grid(self) -> None:
        doc = _make_doc_with_table(cols=3)
        table = doc.tables[0]
        set_table_grid(table, [3000, 3000, 3360])
        xml = table._tbl.xml
        assert 'w:w="3000"' in xml
        assert 'w:w="3360"' in xml

    def test_set_table_borders(self) -> None:
        doc = _make_doc_with_table()
        table = doc.tables[0]
        set_table_borders(table, "D1D5DB")
        xml = table._tbl.xml
        assert "tblBorders" in xml
        assert 'w:color="D1D5DB"' in xml

    def test_set_table_cell_margins(self) -> None:
        doc = _make_doc_with_table()
        table = doc.tables[0]
        set_table_cell_margins(table, top=50, left=100, bottom=50, right=100)
        xml = table._tbl.xml
        assert "tblCellMar" in xml

    def test_set_table_cell_margins_defaults(self) -> None:
        doc = _make_doc_with_table()
        table = doc.tables[0]
        set_table_cell_margins(table)
        xml = table._tbl.xml
        assert 'w:w="40"' in xml
        assert 'w:w="80"' in xml


class TestCellHelpers:
    def test_set_cell_shading(self) -> None:
        doc = _make_doc_with_table()
        cell = doc.tables[0].rows[0].cells[0]
        set_cell_shading(cell, "F3F4F6")
        xml = cell._tc.xml
        assert 'w:fill="F3F4F6"' in xml

    def test_set_cell_width(self) -> None:
        doc = _make_doc_with_table()
        cell = doc.tables[0].rows[0].cells[0]
        set_cell_width(cell, 4680)
        xml = cell._tc.xml
        assert 'w:w="4680"' in xml

    def test_set_cell_vertical_alignment(self) -> None:
        doc = _make_doc_with_table()
        cell = doc.tables[0].rows[0].cells[0]
        set_cell_vertical_alignment(cell)
        xml = cell._tc.xml
        assert 'w:val="center"' in xml

    def test_set_cell_vertical_alignment_top(self) -> None:
        doc = _make_doc_with_table()
        cell = doc.tables[0].rows[0].cells[0]
        set_cell_vertical_alignment(cell, val="top")
        xml = cell._tc.xml
        assert 'w:val="top"' in xml


class TestParagraphHelpers:
    def test_set_paragraph_shading(self) -> None:
        doc = Document()
        p = doc.add_paragraph()
        set_paragraph_shading(p, "F3F4F6")
        xml = p._p.xml
        assert 'w:fill="F3F4F6"' in xml

    def test_set_paragraph_left_border(self) -> None:
        doc = Document()
        p = doc.add_paragraph()
        set_paragraph_left_border(p, "94A3B8")
        xml = p._p.xml
        assert "pBdr" in xml
        assert 'w:color="94A3B8"' in xml

    def test_set_paragraph_bottom_border(self) -> None:
        doc = Document()
        p = doc.add_paragraph()
        set_paragraph_bottom_border(p, "D1D5DB")
        xml = p._p.xml
        assert "pBdr" in xml
        assert 'w:color="D1D5DB"' in xml


class TestRunHelpers:
    def test_set_secondary_font(self) -> None:
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        set_secondary_font(run, "Helvetica Neue")
        xml = run._r.xml
        assert "Helvetica Neue" in xml

    def test_set_secondary_font_idempotent(self) -> None:
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        set_secondary_font(run, "Menlo")
        set_secondary_font(run, "Menlo")
        assert run._r.xml.count("Menlo") >= 1


class TestAddPageField:
    def test_adds_page_field_runs(self) -> None:
        doc = Document()
        p = doc.add_paragraph()
        add_page_field(p, font_name="Aptos", font_size=8.0, color="6B7280")
        xml = p._p.xml
        assert "fldChar" in xml
        assert "PAGE" in xml

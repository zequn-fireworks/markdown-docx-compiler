"""Unit tests for the new-pipeline DocxRenderer and region rendering."""

from __future__ import annotations

import re
from zipfile import ZipFile

from docx.shared import Inches, Pt

from markdown_docx_compiler.backend.docx.doc_renderer import DocxRenderer
from markdown_docx_compiler.models.config import (
    DocumentConfig,
    MarginConfig,
    PageConfig,
    RegionStyle,
)
from markdown_docx_compiler.models.document import (
    BlockMeta,
    Document,
    Heading,
    List,
    ListItem,
    Paragraph,
    Region,
    Table,
    TableCell,
    TextContent,
    TextSpan,
)
from markdown_docx_compiler.models.style import (
    BlockStyle,
    BorderSide,
    BorderStyle,
    FontStyle,
    ListProps,
)
from markdown_docx_compiler.resolve.defaults import DEFAULT_DOCUMENT


class TestDocxRendererSetup:
    def _make_renderer(self, **overrides: object) -> DocxRenderer:
        config = DocumentConfig(
            font=FontStyle(family="Arial", size=11.0, color="222222"),
            page=PageConfig(
                width_inches=8.5,
                margin=MarginConfig(top=1.0, bottom=0.8, left=1.0, right=1.0),
            ),
        )
        return DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=RegionStyle(),
        )

    def test_content_width_inches(self) -> None:
        r = self._make_renderer()
        assert r._content_width_inches == 8.5 - 1.0 - 1.0

    def test_content_width_twips(self) -> None:
        r = self._make_renderer()
        assert r._content_width_twips == int(6.5 * 1440)

    def test_configure_document_sets_margins(self) -> None:
        r = self._make_renderer()
        r._configure_document()
        section = r.document.sections[0]
        assert section.top_margin == Inches(1.0)
        assert section.bottom_margin == Inches(0.8)

    def test_configure_document_sets_normal_font(self) -> None:
        r = self._make_renderer()
        r._configure_document()
        normal = r.document.styles["Normal"]
        assert normal.font.name == "Arial"
        assert normal.font.size == Pt(11.0)

    def test_default_config_uses_defaults(self) -> None:
        r = DocxRenderer(
            config=DocumentConfig(),
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=RegionStyle(),
        )
        assert r._content_width_inches == (
            (DEFAULT_DOCUMENT.page.width_inches or 8.5)
            - (DEFAULT_DOCUMENT.page.margin.left or 1.0)
            - (DEFAULT_DOCUMENT.page.margin.right or 1.0)
        )


class TestDocxRendererRender:
    def test_render_empty_document(self) -> None:
        config = DocumentConfig(font=FontStyle(family="Arial", size=11.0))
        renderer = DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=RegionStyle(),
        )
        ir = Document(metadata={})
        doc = renderer.render(ir, block_styles={})
        assert doc is not None

    def test_render_heading_and_paragraph(self) -> None:
        config = DocumentConfig(font=FontStyle(family="Arial", size=11.0))
        renderer = DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=RegionStyle(),
        )
        ir = Document(
            metadata={},
            body=[
                Heading(level=1, content=[TextSpan("Title")], meta=BlockMeta(index=1)),
                Paragraph(content=[TextSpan("Body text.")], meta=BlockMeta(index=2)),
            ],
        )
        styles = {
            1: BlockStyle(font=FontStyle(size=20.0, bold=True)),
            2: BlockStyle(),
        }
        doc = renderer.render(ir, block_styles=styles)
        texts = [p.text for p in doc.paragraphs]
        assert "Title" in texts
        assert "Body text." in texts

    def test_render_table(self) -> None:
        config = DocumentConfig(font=FontStyle(family="Arial", size=11.0))
        renderer = DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=RegionStyle(),
        )
        ir = Document(
            metadata={},
            body=[
                Table(
                    headers=[TableCell(content=[TextSpan("A")]), TableCell(content=[TextSpan("B")])],
                    rows=[[TableCell(content=[TextSpan("1")]), TableCell(content=[TextSpan("2")])]],
                    alignments=["left", "right"],
                    meta=BlockMeta(index=1),
                ),
            ],
        )
        styles = {1: BlockStyle()}
        doc = renderer.render(ir, block_styles=styles)
        assert len(doc.tables) == 1

    def test_render_list_continuation_paragraph_uses_continue_style(self) -> None:
        config = DocumentConfig(font=FontStyle(family="Arial", size=11.0))
        renderer = DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=RegionStyle(),
        )
        ir = Document(
            metadata={},
            body=[
                List(
                    ordered=True,
                    items=[
                        ListItem(
                            blocks=[
                                Paragraph(content=[TextSpan("Primary item")], meta=BlockMeta(index=1)),
                                Paragraph(content=[TextSpan("Continuation note")], meta=BlockMeta(index=2)),
                            ]
                        )
                    ],
                    meta=BlockMeta(index=3),
                )
            ],
        )
        styles = {
            1: BlockStyle(),
            2: BlockStyle(background="FEF3C7"),
            3: BlockStyle(),
        }

        doc = renderer.render(ir, block_styles=styles)

        assert doc.paragraphs[0].style.name == "List Number"
        assert doc.paragraphs[1].style.name == "List Continue"
        assert doc.paragraphs[1].text == "Continuation note"

    def test_render_separate_ordered_lists_restart_numbering(self, tmp_path: object) -> None:
        from pathlib import Path

        out = Path(str(tmp_path)) / "lists.docx"
        config = DocumentConfig(font=FontStyle(family="Arial", size=11.0))
        renderer = DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=RegionStyle(),
        )
        ir = Document(
            metadata={},
            body=[
                List(ordered=True, items=[ListItem(blocks=[Paragraph(content=[TextSpan("First list item")], meta=BlockMeta(index=1))])]),
                Paragraph(content=[TextSpan("Separator")], meta=BlockMeta(index=2)),
                List(ordered=True, items=[ListItem(blocks=[Paragraph(content=[TextSpan("Second list item")], meta=BlockMeta(index=3))])]),
            ],
        )

        doc = renderer.render(
            ir,
            block_styles={
                1: BlockStyle(),
                2: BlockStyle(),
                3: BlockStyle(),
            },
        )
        doc.save(str(out))

        with ZipFile(out) as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")

        num_ids = re.findall(r"<w:numId w:val=\"(\d+)\"/>", document_xml)
        assert len(num_ids) >= 2
        assert num_ids[0] != num_ids[1]

    def test_render_nested_ordered_lists_use_hierarchical_numbering(self, tmp_path: object) -> None:
        from pathlib import Path

        out = Path(str(tmp_path)) / "nested-lists.docx"
        config = DocumentConfig(font=FontStyle(family="Arial", size=11.0))
        renderer = DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=RegionStyle(),
        )
        ir = Document(
            metadata={},
            body=[
                List(
                    ordered=True,
                    items=[
                        ListItem(
                            blocks=[
                                Paragraph(content=[TextSpan("Parent item")], meta=BlockMeta(index=1)),
                                List(
                                    ordered=True,
                                    items=[
                                        ListItem(blocks=[Paragraph(content=[TextSpan("Child item")], meta=BlockMeta(index=2))]),
                                    ],
                                    meta=BlockMeta(index=3),
                                ),
                            ]
                        )
                    ],
                    meta=BlockMeta(index=4),
                )
            ],
        )

        doc = renderer.render(
            ir,
            block_styles={
                1: BlockStyle(),
                2: BlockStyle(),
                3: BlockStyle(),
                4: BlockStyle(),
            },
        )
        doc.save(str(out))

        with ZipFile(out) as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            numbering_xml = zf.read("word/numbering.xml").decode("utf-8")

        num_entries = re.findall(r"<w:numPr><w:ilvl w:val=\"(\d+)\"/><w:numId w:val=\"(\d+)\"/></w:numPr>", document_xml)
        assert len(num_entries) >= 2
        assert num_entries[0][0] == "0"
        assert num_entries[1][0] == "1"
        assert num_entries[0][1] == num_entries[1][1]
        assert 'w:name w:val="mdc-decimal_hierarchical"' in numbering_xml
        assert 'w:lvlText w:val="%1.%2"' in numbering_xml

    def test_render_alpha_paren_numbering_scheme(self, tmp_path: object) -> None:
        from pathlib import Path

        out = Path(str(tmp_path)) / "alpha-paren.docx"
        config = DocumentConfig(font=FontStyle(family="Arial", size=11.0))
        renderer = DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=RegionStyle(),
        )
        list_style = BlockStyle(list=ListProps(numbering="alpha_paren_hierarchical"))
        ir = Document(
            metadata={},
            body=[
                List(
                    ordered=True,
                    items=[
                        ListItem(
                            blocks=[
                                Paragraph(content=[TextSpan("Parent item")], meta=BlockMeta(index=1)),
                                List(
                                    ordered=True,
                                    items=[
                                        ListItem(blocks=[Paragraph(content=[TextSpan("Child item")], meta=BlockMeta(index=2))]),
                                    ],
                                    meta=BlockMeta(index=3),
                                ),
                            ]
                        )
                    ],
                    meta=BlockMeta(index=4),
                )
            ],
        )

        doc = renderer.render(
            ir,
            block_styles={
                1: BlockStyle(),
                2: BlockStyle(),
                3: list_style,
                4: list_style,
            },
        )
        doc.save(str(out))

        with ZipFile(out) as zf:
            numbering_xml = zf.read("word/numbering.xml").decode("utf-8")

        assert 'w:name w:val="mdc-alpha_paren_hierarchical"' in numbering_xml
        assert 'w:lvlText w:val="%1(%2)"' in numbering_xml

    def test_render_alpha_hierarchical_numbering_scheme(self, tmp_path: object) -> None:
        from pathlib import Path

        out = Path(str(tmp_path)) / "alpha-hierarchical.docx"
        config = DocumentConfig(font=FontStyle(family="Arial", size=11.0))
        renderer = DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=RegionStyle(),
        )
        list_style = BlockStyle(list=ListProps(numbering="alpha_hierarchical"))
        ir = Document(
            metadata={},
            body=[
                List(
                    ordered=True,
                    items=[
                        ListItem(
                            blocks=[
                                Paragraph(content=[TextSpan("Parent item")], meta=BlockMeta(index=1)),
                                List(
                                    ordered=True,
                                    items=[
                                        ListItem(blocks=[Paragraph(content=[TextSpan("Child item")], meta=BlockMeta(index=2))]),
                                    ],
                                    meta=BlockMeta(index=3),
                                ),
                            ]
                        )
                    ],
                    meta=BlockMeta(index=4),
                )
            ],
        )

        doc = renderer.render(
            ir,
            block_styles={
                1: BlockStyle(),
                2: BlockStyle(),
                3: list_style,
                4: list_style,
            },
        )
        doc.save(str(out))

        with ZipFile(out) as zf:
            numbering_xml = zf.read("word/numbering.xml").decode("utf-8")

        assert 'w:name w:val="mdc-alpha_hierarchical"' in numbering_xml
        assert 'w:lvlText w:val="%1.%2"' in numbering_xml

    def test_render_with_page_footer(self, tmp_path: object) -> None:
        config = DocumentConfig(font=FontStyle(family="Arial", size=11.0))
        footer_style = RegionStyle(
            font=FontStyle(size=8.0, color="888888"),
            border=BorderStyle(top=BorderSide(color="CCCCCC")),
        )
        renderer = DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=footer_style,
            doc_header_style=RegionStyle(),
        )
        region = Region(
            center=[TextContent(content=[TextSpan("Page {page}")])],
        )
        ir = Document(metadata={}, page_footer=region)
        doc = renderer.render(ir, block_styles={})
        assert doc is not None

    def test_render_with_doc_header(self) -> None:
        config = DocumentConfig(font=FontStyle(family="Arial", size=11.0))
        header_style = RegionStyle(font=FontStyle(size=9.0, color="666666"))
        renderer = DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=header_style,
        )
        region = Region(
            left=[TextContent(content=[TextSpan("Company Inc.")])],
            right=[TextContent(content=[TextSpan("2026-01-01")])],
        )
        ir = Document(metadata={}, doc_header=region)
        doc = renderer.render(ir, block_styles={})
        assert len(doc.tables) == 1


class TestDocxRendererSave:
    def test_save_produces_valid_docx(self, tmp_path: object) -> None:
        from pathlib import Path

        out = Path(str(tmp_path)) / "test.docx"
        config = DocumentConfig(font=FontStyle(family="Arial", size=11.0))
        renderer = DocxRenderer(
            config=config,
            page_header_style=RegionStyle(),
            page_footer_style=RegionStyle(),
            doc_header_style=RegionStyle(),
        )
        ir = Document(
            metadata={},
            body=[Paragraph(content=[TextSpan("Hello")], meta=BlockMeta(index=1))],
        )
        doc = renderer.render(ir, block_styles={1: BlockStyle()})
        doc.save(str(out))

        assert out.exists()
        with ZipFile(out) as zf:
            names = zf.namelist()
            assert "word/document.xml" in names
